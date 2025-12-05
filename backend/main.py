# main.py
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pathlib import Path
from collections import Counter
import re, os, string, math, asyncio, time, json, requests
from typing import List, Dict, Any
from datetime import datetime

from colorama import Fore, Style, init as colorama_init
colorama_init(autoreset=True)

# CONFIG & PARAMETER
OLLAMA_API_URL = os.getenv("OLLAMA_API_URL", "http://localhost:11434/api/generate")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1:70b")

MAX_CONCURRENCY = 6
MAX_RETRIES = 4
RETRY_BASE_DELAY = 0.8
OLLAMA_TIMEOUT = 600
MAX_INPUT_CHARS = 50000

# PDF TEXT EXTRACTION
def read_pdf_text(path: str) -> str:
    text = ""
    try:
        import fitz
        doc = fitz.open(path)
        text = "\n".join([p.get_text() for p in doc])
        text = clean_text(text)
        if _enough_text(text):
            return text
    except:
        pass

    try:
        from pdfminer.high_level import extract_text
        text = clean_text(extract_text(path))
        if _enough_text(text):
            return text
    except:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        raw = "\n".join([p.extract_text() or "" for p in reader.pages])
        text = clean_text(raw)
        if _enough_text(text):
            return text
    except:
        pass

    try:
        from pdf2image import convert_from_path
        import pytesseract

        pages = convert_from_path(path, dpi=350)

        ocr_text = ""
        for i, pg in enumerate(pages):
            img_text = pytesseract.image_to_string(pg, lang="ind+eng")
            # Prioritaskan halaman yang mengandung kata 'BIODATA'
            if "BIODATA" in img_text.upper() or i > len(pages)-3:  
                ocr_text += "\n--- OCR PAGE %d ---\n" % (i+1)
                ocr_text += img_text + "\n"

        if len(text) < 200 or len(ocr_text) > len(text)*0.3:
            return clean_text(text + "\n\n" + ocr_text)
    except:
        pass

    return clean_text(text)

def _enough_text(text, min_chars=200):
    return len(text.strip()) >= min_chars

BLACKBOX = ["■","□","▯","█","�"]
def clean_text(text):
    if not text:
        return ""
    for b in BLACKBOX:
        text = text.replace(b, "")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"Gambar\s*\d+(\.\d+)*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"Tabel\s*\d+(\.\d+)*", "", text, flags=re.IGNORECASE)
    return text.strip()

def clean_reference_noise(text):
    text = re.sub(r"http\S+|www\S+", "", text)
    text = re.sub(r"\([A-Za-z][^()]{0,40}\d{4}\)", "", text)
    text = re.sub(r"[A-Za-z]+,\s*\d{4}", "", text)
    text = re.sub(r"([A-Za-z]+\s*,){2,}.*", "", text)
    text = re.sub(r"(Universitas|Fakultas|Program Studi|Jurusan|Departemen).*", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

def remove_duplicate_paragraphs(text: str) -> str:
    """
    Menghapus paragraf atau kalimat yang muncul dua kali (duplikasi PDF).
    Cocok untuk PDF skripsi yang layer text-nya double.
    """
    if not text:
        return text

    paras = [p.strip() for p in text.split("\n") if p.strip()]
    unique = []
    seen = set()

    for p in paras:
        key = p[:120].lower()  # fingerprint pendek
        if key not in seen:
            seen.add(key)
            unique.append(p)

    return "\n".join(unique)
    
def remove_bab_intro_paragraph(text: str) -> str:
    """
    Menghapus paragraf pembuka seperti:
    - 'Bab ini menguraikan...'
    - 'Bab X membahas...'
    - 'Bab ini akan menjelaskan...'
    dan membuang paragraf duplikat otomatis.
    """
    if not text:
        return text

    # buang paragraf pembuka deskriptif
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    clean_paragraphs = []
    
    intro_pattern = re.compile(
        r"^\s*(bab\s*(i|ii|iii|iv|v|\d+)?\s*(ini)?\s*(akan\s+)?"
        r"(membahas|menguraikan|menjelaskan|memaparkan|menjabarkan))",
        flags=re.IGNORECASE
    )

    for p in paragraphs:
        if intro_pattern.search(p):
            continue
        clean_paragraphs.append(p)

    # hilangkan duplikasi paragraf yang sama
    final_unique = []
    seen = set()
    for p in clean_paragraphs:
        key = p[:80].lower()
        if key not in seen:
            seen.add(key)
            final_unique.append(p)

    return "\n".join(final_unique)

def remove_subbab(text: str) -> str:
    # Hilangkan penomoran subbab (3.1, 3.2.1, dst.)
    return re.sub(r"\b\d+\.\d+(\.\d+)*\b", "", text)

# SPLIT BAB
def split_by_bab(text: str):
    # Hapus daftar isi, daftar tabel/gambar, daftar pustaka, lampiran
    text = re.sub(r"DAFTAR\s+ISI.*?(?=BAB\s+[IVXLCDM1]\b)", "", text,
                  flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"DAFTAR\s+(GAMBAR|TABEL).*?(?=BAB\s+[IVXLCDM1]\b)", "", text,
                  flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"DAFTAR PUSTAKA.*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"LAMPIRAN.*", "", text, flags=re.IGNORECASE)

    # Hilangkan titik-titik (.......) yang sering mengacaukan header
    text = re.sub(r"^.*\.{5,}.*$", "", text, flags=re.MULTILINE)

    # Buang halaman romawi (i, ii, iii) yang berdiri sendiri
    text = re.sub(r"(?m)^\s*[ivxlcdm]+\s*$", "", text, flags=re.IGNORECASE)

    # Mulai dari BAB I/1 (abaikan halaman depan)
    m = re.search(r"(BAB\s+(?:I|II|III|IV|V|1|2|3|4|5)\b.*)", text,
                  flags=re.IGNORECASE | re.DOTALL)
    if m:
        text = m.group(1)

    # Pecah berdasarkan heading BAB ketat
    parts = re.split(
        r"(?mi)^(BAB[\s]+(?:I|II|III|IV|V|VI|VII|VIII|IX|X|\b1\b|\b2\b|\b3\b|\b4\b|\b5\b))\s*$",
        text
    )

    result = []
    for i in range(1, len(parts), 2):
        judul = parts[i].strip()
        isi = parts[i+1].strip() if (i+1 < len(parts)) else ""

        # Filter keras:
        # Jangan pernah menangkap "Subbab 4.1"
        if re.search(r"\b\d+\.\d+(\.\d+)*\b", judul):
            continue

        # Minimal konten banget (kalau terlalu sedikit, ignore)
        if len(isi) < 200:
            continue

        # Bersihkan internal 'bab 4.1.2 ...'
        isi = re.sub(r"(?im)^\s*bab\s+\d+(\.\d+)*.*$", "", isi)
        isi = re.sub(r"(?im)^\s*\d+(\.\d+)+.*$", "", isi)
        # Bersihkan header subbab 3.1 / 4.2 / 3.1.1
        isi = re.sub(r"(?m)^\s*\d+(\.\d+){1,3}.*$", "", isi)
        isi = re.sub(r"(?i)\bsubbab\s+\d+(\.\d+)*\b.*", "", isi)

        if len(isi) < 200:
            continue

        result.append({"judul": judul, "isi": isi})

    return result

    # Kata kunci teknis yang menaikkan skor (bahasa Indonesia + simbol)
    KEYWORDS = [
        "metode","sintesis","represipitasi","psa","karakteris","karakteriza",
        "imobilis","imobiliza","µpad","μpad","nanokristal","bhb","triptamin",
        "imagej","uv","fluores","emisi","analisis","hasil","pembahas","validasi",
        "dispersi","konsentrasi","kecap","sampling","pengujian","selektivitas"
    ]

    def score_text(t: str) -> int:
        s = 0
        low_t = t.lower()
        # dasar: panjang
        s += min(len(low_t), 20000)
        # kata kunci
        key_count = sum(1 for k in KEYWORDS if k in low_t)
        s += key_count * 800
        # jumlah kalimat berguna
        sent_count = len(re.findall(r'[\.!?]', low_t))
        s += min(sent_count, 50) * 50
        # angka/ukur (adanya angka biasanya tanda data atau parameter)
        if re.search(r'\d', low_t):
            s += 500
        # jika ada banyak istilah ilmiah (huruf panjang kata)
        long_word_count = sum(1 for w in re.findall(r'\w+', low_t) if len(w) > 6)
        s += min(long_word_count, 200) * 5
        # penalti jika hanya frasa meta seperti "Bab ini membahas" tanpa kata kunci
        if re.search(r'\bbab\s+\w+\s+membahas', low_t) and key_count == 0 and len(low_t) < 1000:
            s -= 10000
        return s

    # Kelompokkan kandidat berdasarkan judul (BAB I, BAB II, ...)
    groups = {}
    for c in candidates:
        key = re.sub(r'\s+', ' ', c["judul"].upper().strip())
        groups.setdefault(key, []).append(c)

    # Pilih kandidat terbaik per grup (skor tertinggi), simpan pos aslinya
    chosen = []
    for key, items in groups.items():
        best = max(items, key=lambda it: score_text(it["isi"]))
        best["score"] = score_text(best["isi"])
        chosen.append(best)

    # Urutkan berdasarkan posisi terawal kemunculan di dokumen
    chosen.sort(key=lambda x: x["pos"])

    # Final cleaning: buang yang sangat pendek dan tidak informatif
    final = []
    for ch in chosen:
        isi_bersih = re.sub(r'\s+', ' ', ch["isi"]).strip()
        # jika sangat pendek dan tidak mengandung kata kunci penting, skip
        if len(isi_bersih) < 400 and all(k not in isi_bersih.lower() for k in KEYWORDS):
            print(f"{Fore.YELLOW}[FILTER]{Style.RESET_ALL} Menghapus {ch['judul']} (terlalu pendek/tidak teknis).")
            continue
        # Hapus paragraf intro “Bab ini membahas …”
        isi_final = remove_bab_intro_paragraph(ch["isi"])
        final.append({"judul": ch["judul"], "isi": isi_final})

    return final

# UTIL: Tokenisasi & Ringkasan Ekstraktif Lokal
STOPWORDS = set("yang dan di ke dari untuk pada adalah dengan dalam ini itu serta juga tidak dapat atau oleh bagi agar sudah akan para sebagai tersebut karena maka sehingga terhadap serta olehnya".split())
PUNCT = str.maketrans("", "", string.punctuation)

def tokenize(text: str):
    return [w for w in text.lower().translate(PUNCT).split() if w not in STOPWORDS and len(w) > 2]

def split_sentences(text: str):
    sents = re.split(r"(?<=[\.\?\!])\s+(?=[A-Za-z0-9])", text.strip())
    return [s.strip() for s in sents if s.strip()]

def summarize_text_extractive(text: str, max_sent: int = 8) -> str:
    sents = split_sentences(text)
    if not sents: return ""
    sent_tokens = [tokenize(s) for s in sents]
    df = Counter()
    for t in sent_tokens: df.update(set(t))
    N = len(sents)
    scores = []
    for i, toks in enumerate(sent_tokens):
        score = sum((cnt / (1 + len(toks))) * (math.log((N + 1) / (1 + df[w])) + 1)
                    for w, cnt in Counter(toks).items())
        if i < max(3, int(N * 0.1)): score *= 1.15
        scores.append(score)
    top_idx = sorted(range(N), key=lambda i: scores[i], reverse=True)[:max_sent]
    return " ".join([sents[i] for i in sorted(top_idx)])

# PROMPT TEMPLATE
SUM_PROMPT_TEMPLATE = (
    "Anda bertugas membuat dua jenis ringkasan dari satu BAB skripsi.\n\n"
    "1) TLDR (sangat singkat):\n"
    "- Hanya 1 kalimat.\n"
    "- Harus berbeda total dari ringkasan lengkap.\n"
    "- Merangkum inti BAB dalam kalimat paling ringkas.\n"
    "- Tidak boleh mengulang kalimat atau pola bahasa dari ringkasan lengkap.\n\n"
    "2) Ringkasan Lengkap (1–2 paragraf):\n"
    "- Sesuai fungsi BAB:\n"
    "  • BAB I → latar belakang, masalah, tujuan, ruang lingkup\n"
    "  • BAB II → teori, konsep utama, penelitian terdahulu\n"
    "  • BAB III → metode, alat & bahan, alur penelitian\n"
    "  • BAB IV → hasil, temuan, pembahasan\n"
    "  • BAB V → kesimpulan & saran\n"
    "- Bahasa ilmiah, padat, tidak repetitif.\n"
    "Aturan tambahan:\n"
    "- DILARANG membuat heading baru seperti "bab 3.1.1.\n"
    "- Jangan mengulang kalimat dari teks asli.\n"
    "- Jangan membuat 2 paragraf yang maknanya sama.\n"
    "- Hilangkan teks meta seperti 'Bab ini membahas...' dan referensi.\n"
    "- TLDR dan ringkasan lengkap harus berbeda total.\n\n"
    "Format output WAJIB:\n"
    "TLDR:\n"
    "<isi tldr>\n\n"
    "RINGKASAN:\n"
    "<isi ringkasan>\n\n"
    "TEKS SUMBER:\n\"\"\"{content}\"\"\"\n"
)

# OLLAMA CLIENT DENGAN LOG WARNA
def _ollama_generate(prompt: str) -> str:
    try:
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0,
                "top_p": 1,
                "top_k": 1,
                "repeat_penalty": 1.1
            }
        }

        resp = requests.post(
            OLLAMA_API_URL,
            json=payload,
            timeout=OLLAMA_TIMEOUT
        )
        resp.raise_for_status()

        data = resp.json()
        return (data.get("response") or "").strip()

    except Exception as e:
        print(f"{Fore.RED}[OLLAMA ERROR]{Style.RESET_ALL} {e}")
        return ""

async def ollama_summarize_async(content: str, semaphore: asyncio.Semaphore) -> str:
    prompt = SUM_PROMPT_TEMPLATE.format(content=content)
    attempt = 0
    while True:
        attempt += 1
        try:
            async with semaphore:
                start = time.perf_counter()
                result = await asyncio.to_thread(_ollama_generate, prompt)
                elapsed = time.perf_counter() - start
            if result:
                print(f"{Fore.GREEN}[OLLAMA OK]{Style.RESET_ALL} Ringkasan selesai dalam {elapsed:.1f}s (percobaan ke-{attempt})")
                return result
            raise RuntimeError("Empty response from Ollama.")
        except Exception as e:
            if attempt < MAX_RETRIES:
                delay = min(RETRY_BASE_DELAY * (2 ** (attempt - 1)), 8.0)
                print(f"{Fore.YELLOW}[RETRY]{Style.RESET_ALL} Ollama gagal (percobaan ke-{attempt}): {e}. Menunggu {delay:.1f}s...")
                time.sleep(delay)
                continue
            else:
                print(f"{Fore.RED}[FALLBACK]{Style.RESET_ALL} Semua percobaan gagal, pakai ringkasan lokal (TF-IDF).")
                return summarize_text_extractive(content, max_sent=7)

# RINGKAS PDF PER BAB
def compress_for_prompt(text: str, max_chars: int = MAX_INPUT_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    base_k = 10 + min(4, len(text) // 20000)
    extract = summarize_text_extractive(text, max_sent=base_k)
    return extract[:max_chars]

async def summarize_sections_parallel(sections: List[Dict[str, str]]) -> List[Dict[str, Any]]:
    semaphore = asyncio.Semaphore(MAX_CONCURRENCY)

    def extract_parts(output: str):
        tldr = ""
        ringk = ""

        tldr_match = re.search(
            r"TLDR:\s*(.*?)\s*RINGKASAN:",
            output,
            flags=re.DOTALL | re.IGNORECASE
        )
        if tldr_match:
            tldr = tldr_match.group(1).strip()

        ringk_match = re.search(
            r"RINGKASAN:\s*(.*)",
            output,
            flags=re.DOTALL | re.IGNORECASE
        )
        if ringk_match:
            ringk = ringk_match.group(1).strip()

        return tldr, ringk

    async def _process(sec):
        teks = (sec.get("isi") or "").strip()
        if len(teks) < 80:
            return {"judul": sec["judul"], "ringkasan_bab": "", "tldr": ""}

        # BERSIHKAN TEKS
        paragraphs = [p.strip() for p in teks.split("\n") if len(p.strip()) > 40]
        isi_bersih = remove_bab_intro_paragraph("\n".join(paragraphs))
        isi_bersih = clean_reference_noise(isi_bersih)

        # Hapus rujukan subbab di dalam kalimat
        isi_bersih = re.sub(r"(?im)^.*subbab\s+\d+(\.\d+)*.*$", "", isi_bersih)
        isi_bersih = re.sub(r"(?im)^.*bab\s+\d+(\.\d+)*.*$", "", isi_bersih)

        # Hapus baris header “3.1 / 3.1.1 / 4.2.3”
        isi_bersih = re.sub(r"(?m)^\s*(?:\d+\.){1,3}.*$", "", isi_bersih)

        isi_kompres = compress_for_prompt(isi_bersih, MAX_INPUT_CHARS)

        # RINGKASAN VIA OLLAMA
        llm_output = await ollama_summarize_async(isi_kompres, semaphore)
        llm_output = llm_output.strip()

        llm_tldr, llm_ringkas = extract_parts(llm_output)

        # NORMALISASI RINGKASAN
        out_paras = [p.strip() for p in llm_ringkas.split("\n") if p.strip()]
        dedup = []
        seen = set()
        for p in out_paras:
            key = re.sub(r"\s+", " ", p.lower())[:90]
            if key not in seen:
                seen.add(key)
                dedup.append(p)
        final_summary = "\n\n".join(dedup)

        # TLDR
        judul_bab = sec["judul"].lower()

        if "bab i" in judul_bab or "bab 1" in judul_bab:
            tldr_prompt = f"""
Buat TLDR BAB I dengan format berikut:

Latar belakang: <isi satu kalimat>.
Rumusan masalah: <isi satu kalimat>.
Tujuan: <isi satu kalimat>.
Manfaat: <isi satu kalimat>.

Aturan:
- Tidak menyalin kalimat dari ringkasan.
- Hanya isi setelah titik dua.

Ringkasan lengkap:
{final_summary}

Isi TLDR BAB I:
"""
        elif "bab ii" in judul_bab or "bab 2" in judul_bab:
            tldr_prompt = f"""
Buat TLDR BAB II dalam format:

Teori utama: ...
Konsep kunci: ...
Penelitian terdahulu: ...

Ringkasan lengkap:
{final_summary}

Isi TLDR BAB II:
"""
        elif "bab iii" in judul_bab or "bab 3" in judul_bab:
            tldr_prompt = f"""
Buat TLDR BAB III dalam format:

Desain penelitian: ...
Metode: ...
Alat/bahan: ...
Teknik analisis: ...

Ringkasan lengkap:
{final_summary}

Isi TLDR BAB III:
"""
        elif "bab iv" in judul_bab or "bab 4" in judul_bab:
            tldr_prompt = f"""
Buat TLDR BAB IV dalam format:

Hasil utama: ...
Interpretasi: ...
Pembahasan: ...

Ringkasan lengkap:
{final_summary}

Isi TLDR BAB IV:
"""
        elif "bab v" in judul_bab or "bab 5" in judul_bab:
            tldr_prompt = f"""
Buat TLDR BAB V dalam format:

Kesimpulan: ...
Saran: ...

Ringkasan lengkap:
{final_summary}

Isi TLDR BAB V:
"""
        else:
            tldr_prompt = f"""
Ringkas menjadi 3–4 poin inti:

{final_summary}
"""

        # GENERATE TLDR FINAL
        tldr_final = await asyncio.to_thread(_ollama_generate, tldr_prompt)
        tldr_final = (tldr_final or "").strip()
        
        return {
            "judul": sec["judul"],
            "ringkasan_bab": final_summary,
            "tldr": tldr_final
        }

    return await asyncio.gather(
        *[asyncio.create_task(_process(sec)) for sec in sections]
    )

def detect_non_thesis(text: str) -> bool:
    if not text or len(text) < 1000: return True
    t = text.lower()
    bab_count = len(re.findall(r"\b(bab\s+(i|ii|iii|iv|v|1|2|3|4|5))\b", t))
    if bab_count < 2: return True
    keywords = ["pendahuluan","tinjauan pustaka","metodologi","hasil","kesimpulan","rumusan masalah","tujuan"]
    if sum(1 for kw in keywords if kw in t) < 3: return True
    if any(x in t for x in ["invoice","laporan keuangan","brosur","sertifikat"]): return True
    return False

async def summarize_pdf_per_bab(path: str):
    raw = read_pdf_text(path)
    if not raw.strip():
        return {
            "file": os.path.basename(path),
            "sections": [],
            "note": "File kosong atau tidak dapat dibaca.",
            "raw_text": "",
            "bab_sections": []
        }

    raw_clean = remove_duplicate_paragraphs(raw)
    raw_clean = clean_reference_noise(raw_clean)
    raw_clean = remove_subbab(raw_clean)

    bab_sections_raw = split_by_bab(text=raw_clean)

    if not bab_sections_raw:
        bab_sections_raw = [{"judul": "BAB I", "isi": raw_clean}]

    bab_sections_clean = []
    for sec in bab_sections_raw:
        cleaned = sec["isi"]

        # Buang subbab, 3.1.1, 4.2.3, dsb
        cleaned = re.sub(r"(?im)\bsubbab\s*\d+(\.\d+)*\b.*", "", cleaned)
        cleaned = re.sub(r"(?m)^\s*(?:bab\s*)?\d+(\.\d+)+.*$", "", cleaned)

        cleaned = cleaned.strip()

        bab_sections_clean.append({
            "judul": sec["judul"],
            "isi": cleaned
        })

    results = await summarize_sections_parallel(bab_sections_raw)

    return {
        "file": os.path.basename(path),
        "sections": results,          # hasil ringkasan
        "raw_text": raw_clean,        # full teks
        "bab_sections": bab_sections_clean  # untuk UI
    }

# EKSPOR DOCX & PDF
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def export_all(data, out_docx, out_pdf):
    doc = Document()
    doc.add_heading("Ringkasan Per Bab (Ollama)", 0)
    doc.add_paragraph(f"File: {data['file']}")
    for sec in data["sections"]:
        doc.add_heading(sec["judul"], level=1)
        doc.add_paragraph(sec["ringkasan_bab"] or "")
    doc.save(out_docx)

    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(out_pdf, pagesize=A4)
    elements = [Paragraph("Ringkasan Per Bab (Ollama)", styles['Title']),
                Paragraph(f"File: {data['file']}", styles['Normal']),
                Spacer(1, 12)]
    for sec in data["sections"]:
        elements.append(Paragraph(sec["judul"], styles['Heading2']))
        elements.append(Paragraph(sec["ringkasan_bab"] or "", styles['Normal']))
        elements.append(Spacer(1, 12))
    pdf.build(elements)

# qna
def load_doc_context(file_path: Path):
    """
    Mengambil teks full + struktur BAB dari file pendamping.
    Jika belum ada, akan dibuat dari PDF.
    """
    full_path = file_path.with_suffix(".full.txt")
    bab_path = file_path.with_suffix(".bab.json")

    raw_text = ""
    bab_sections: List[Dict[str, str]] = []

    if full_path.exists():
        try:
            raw_text = full_path.read_text(encoding="utf-8", errors="ignore")
        except:
            raw_text = ""

    if bab_path.exists():
        try:
            with open(bab_path, "r", encoding="utf-8") as f:
                bab_sections = json.load(f)
        except:
            bab_sections = []

    # fallback: baca ulang dari PDF
    if not raw_text or not bab_sections:
        pdf_raw = read_pdf_text(str(file_path))
        pdf_clean = remove_duplicate_paragraphs(pdf_raw)
        pdf_clean = clean_reference_noise(pdf_clean)
        pdf_clean = remove_subbab(pdf_clean)

        raw_text = raw_text or pdf_clean
        if not bab_sections:
            bab_sections = split_by_bab(pdf_clean) or [{"judul": "BAB I", "isi": pdf_clean}]

        # simpan lagi supaya next call lebih cepat
        try:
            full_path.write_text(raw_text, encoding="utf-8")
            with open(bab_path, "w", encoding="utf-8") as f:
                json.dump(bab_sections, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"{Fore.YELLOW}[WARN]{Style.RESET_ALL} Gagal menyimpan konteks (lazy): {e}")

    return raw_text, bab_sections


def pick_relevant_babs(question: str, bab_sections: List[Dict[str, str]], top_k: int = 3):
    """
    Pilih BAB yang paling relevan dengan pertanyaan
    pakai overlap token sederhana (IR hemat biaya).
    """
    if not bab_sections:
        return []

    q_tokens = set(tokenize(question))
    scored = []
    for sec in bab_sections:
        isi = sec.get("isi", "") or ""
        judul = sec.get("judul", "") or ""
        teks_tokens = set(tokenize(isi))
        overlap = len(q_tokens & teks_tokens)
        # bonus kalau kata pertanyaan muncul di judul
        judul_low = judul.lower()
        for t in q_tokens:
            if t in judul_low:
                overlap += 2
        scored.append((overlap, sec))

    scored.sort(key=lambda x: x[0], reverse=True)
    chosen = [sec for score, sec in scored[:top_k] if score > 0]

    # kalau tidak ada yang match, ambil BAB pertama saja
    if not chosen:
        chosen = bab_sections[:1]
    return chosen

# FASTAPI APP
app = FastAPI(title="DocuSum AI (Ollama)", version="9.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
BASE_URL = "https://docusum.onrender.com"

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Hanya file PDF diperbolehkan.")
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as f:
        f.write(await file.read())
    
    # ---- Antivirus Scan ----
    scan_result = os.popen(f"clamscan --infected --no-summary {str(file_path)}").read()

    if "FOUND" in scan_result:
        os.remove(file_path)
        raise HTTPException(
            status_code=400,
            detail="File terdeteksi virus dan telah ditolak."
        )

    print("[SCAN] ClamAV result:", scan_result)

    try:
        hasil = await summarize_pdf_per_bab(str(file_path))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal meringkas: {e}")
            # simpan konteks untuk QnA / chat
    raw_text = hasil.get("raw_text", "")
    bab_sections = hasil.get("bab_sections", [])

    try:
        if raw_text:
            with open(file_path.with_suffix(".full.txt"), "w", encoding="utf-8") as f:
                f.write(raw_text)
        if bab_sections:
            with open(file_path.with_suffix(".bab.json"), "w", encoding="utf-8") as f:
                json.dump(bab_sections, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"{Fore.YELLOW}[WARN]{Style.RESET_ALL} Gagal menyimpan konteks QnA: {e}")
    if hasil.get("note"):
        return {"success": False, "message": hasil["note"], "file": hasil["file"]}
    docx_path = str(file_path.with_suffix(".docx"))
    pdf_path = str(file_path.with_suffix(".summary.pdf"))
    export_all(hasil, docx_path, pdf_path)
    hasil["download_docx"] = f"{BASE_URL}/api/download/{Path(docx_path).name}"
    hasil["download_pdf"] = f"{BASE_URL}/api/download/{Path(pdf_path).name}"
    return {
        "success": True,
        "file": file.filename,  # ⬅ ini penting!
        "data": hasil
    }
    
@app.get("/api/download/{filename}")
async def download_file(filename: str):
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File tidak ditemukan")
    return FileResponse(file_path, filename=filename)

# prompt tanya
@app.post("/api/ask")
async def ask_about_file(data: Dict[str, Any]):
    question = (data.get("question") or "").strip()
    file_name = (data.get("file") or "").strip()

    if not question:
        raise HTTPException(status_code=400, detail="Pertanyaan tidak boleh kosong.")
    if not file_name:
        raise HTTPException(status_code=400, detail="Nama file wajib diisi.")

    file_path = UPLOAD_DIR / file_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File tidak ditemukan.")

    raw_text, bab_sections = load_doc_context(file_path)
    if not raw_text.strip():
        return {"answer": "Dokumen tidak terbaca atau kosong."}

    MAX_QA_CHARS = 400_000

    # ================== CONTEXT BUILDER ==================
    if bab_sections:
        relevant = pick_relevant_babs(question, bab_sections, top_k=3)

        parts = []
        per_sec_limit = MAX_QA_CHARS // max(len(relevant), 1)

        for sec in relevant:
            isi = sec.get("isi", "")
            if len(isi) > per_sec_limit:
                isi = isi[:per_sec_limit]
            parts.append(f"{sec['judul']}\n{isi}")

        context = "\n\n".join(parts)

        context += "\n\n--- HALAMAN AWAL ---\n"
        context += raw_text[:5000]

        context += "\n\n--- HALAMAN AKHIR ---\n"
        context += raw_text[-4000:]

    else:
        context = raw_text[:MAX_QA_CHARS]
    
    prompt = f"""
Anda adalah AI akademik yang menjawab pertanyaan berdasarkan isi skripsi di bawah ini.

=== DOKUMEN SUMBER (TEKS PENUH) ===
{context}

=== PERTANYAAN PENGGUNA ===
{question}

Instruksi jawaban:
- Jawab hanya berdasarkan dokumen sumber di atas.
- Jika tampak berasal dari BAB tertentu, sebutkan BAB tersebut di jawaban.
- Jika informasi tidak ditemukan di dokumen, jawab dengan jujur bahwa informasi tidak ada.
- Jawab dengan bahasa Indonesia yang natural, jelas, dan padat (seperti ChatGPT).

Jawaban:
"""
    jawaban = (await asyncio.to_thread(_ollama_generate, prompt)) or ""

    return {"answer": jawaban.strip()}

@app.post("/api/search")
async def search_in_file(data: Dict[str, str]):
    """
    Body JSON:
    {
        "query": "kualitas pelayanan",
        "file": "nama.pdf"
    }
    """
    query = (data.get("query") or "").strip()
    file_name = (data.get("file") or "").strip()

    if not query:
        raise HTTPException(status_code=400, detail="Kata kunci tidak boleh kosong.")
    if not file_name:
        raise HTTPException(status_code=400, detail="Nama file wajib diisi.")

    file_path = UPLOAD_DIR / file_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File tidak ditemukan.")

    _, bab_sections = load_doc_context(file_path)
    q_low = query.lower()

    hits = []
    for sec in bab_sections:
        judul = sec.get("judul", "BAB ?")
        isi = sec.get("isi", "") or ""
        for para in isi.split("\n"):
            if q_low in para.lower():
                snippet = para.strip()
                if snippet:
                    hits.append({
                        "bab": judul,
                        "snippet": snippet[:400]
                    })

    return {
        "query": query,
        "file": file_name,
        "results": hits[:20]   # batasi 20 hasil
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)



